"""
docx_exporter.py
Export review questions and author responses to Word documents using a format config JSON.
"""

from __future__ import annotations

import json
import os
from datetime import datetime
from typing import Any, cast

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_REVIEW_REQUIRED_KEYS = ("reference location", "reference text", "issue", "detail")
_RESPONSE_REQUIRED_KEYS = ("problem", "responde")
_CODE_FENCE_PREFIXES = ("```json", "```")

_DEFAULT_FORMAT: dict[str, Any] = {
    "document": {
        "default_font": {
            "latin": "Times New Roman",
            "east_asia": "宋体",
            "size": 12,
        },
        "margins_cm": {
            "top": 2.54,
            "bottom": 2.54,
            "left": 3.18,
            "right": 3.18,
        },
    },
    "paragraph": {
        "line_spacing": 1.5,
        "space_after_pt": 6,
        "first_line_indent_chars": 2.0,
    },
    "title": {
        "font": {"latin": "Times New Roman", "east_asia": "黑体", "size": 16, "bold": True},
        "align": "center",
        "space_after_pt": 12,
    },
    "section_heading": {
        "font": {"latin": "Times New Roman", "east_asia": "黑体", "size": 14, "bold": True},
        "space_before_pt": 12,
        "space_after_pt": 6,
    },
    "metadata": {
        "font": {"latin": "Times New Roman", "east_asia": "宋体", "size": 12},
    },
    "review": {
        "number_font": {"latin": "Times New Roman", "east_asia": "黑体", "size": 12, "bold": True},
        "label_font": {"latin": "Times New Roman", "east_asia": "黑体", "size": 12, "bold": True},
        "body_font": {"latin": "Times New Roman", "east_asia": "宋体", "size": 12},
        "left_indent_cm": 0.74,
        "reference_color_rgb": "8B0000",
        "labels": {
            "reference_location": "引用位置：",
            "reference_text": "原文：",
            "issue": "问题：",
            "detail": "详情：",
        },
    },
    "response": {
        "title_font": {"latin": "Times New Roman", "east_asia": "黑体", "size": 12, "bold": True},
        "label_font": {"latin": "Times New Roman", "east_asia": "黑体", "size": 12, "bold": True},
        "body_font": {"latin": "Times New Roman", "east_asia": "宋体", "size": 12},
        "left_indent_cm": 0.74,
        "labels": {
            "problem": "审稿问题总结：",
            "responde": "回应与解决方法：",
        },
    },
    "footer": {
        "font": {"latin": "Times New Roman", "east_asia": "宋体", "size": 10},
        "color_rgb": "666666",
        "space_before_pt": 18,
    },
}


def _deep_merge(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    merged = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = _deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def _load_format_config(format_path: str | None) -> dict[str, Any]:
    if format_path is None:
        format_path = os.path.join(os.path.dirname(__file__), "format_zh_academic.json")

    if not os.path.exists(format_path):
        raise RuntimeError(f"未找到格式配置文件: {format_path}")

    try:
        with open(format_path, "r", encoding="utf-8") as fh:
            user_config = json.load(fh)
    except OSError as exc:
        raise RuntimeError(f"读取格式配置文件失败: {format_path}") from exc
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"格式配置 JSON 无效: {format_path}") from exc

    if not isinstance(user_config, dict):
        raise RuntimeError("格式配置文件顶层必须是对象")

    return _deep_merge(_DEFAULT_FORMAT, user_config)


def _set_run_font(run, font_config: dict[str, Any], bold: bool | None = None):
    latin = str(font_config.get("latin", "Times New Roman"))
    east_asia = str(font_config.get("east_asia", "宋体"))
    size = float(font_config.get("size", 12))

    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bool(font_config.get("bold", False)) if bold is None else bold

    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:eastAsia"), east_asia)


def _set_rgb_color(run, rgb_hex: str | None):
    if not rgb_hex:
        return
    cleaned = rgb_hex.strip().lstrip("#")
    if len(cleaned) != 6:
        return
    try:
        run.font.color.rgb = RGBColor(int(cleaned[0:2], 16), int(cleaned[2:4], 16), int(cleaned[4:6], 16))
    except ValueError:
        return


def _set_paragraph_format(paragraph, paragraph_config: dict[str, Any], first_line_chars: float | None = None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = float(paragraph_config.get("line_spacing", 1.5))
    fmt.space_after = Pt(float(paragraph_config.get("space_after_pt", 6)))

    if first_line_chars is None:
        first_line_chars = paragraph_config.get("first_line_indent_chars", 2.0)
    if first_line_chars is not None:
        fmt.first_line_indent = Cm(0.74 * float(first_line_chars))


def _set_page_layout(document: DocxDocument, margins: dict[str, Any]):
    section = document.sections[0]
    section.top_margin = Cm(float(margins.get("top", 2.54)))
    section.bottom_margin = Cm(float(margins.get("bottom", 2.54)))
    section.left_margin = Cm(float(margins.get("left", 3.18)))
    section.right_margin = Cm(float(margins.get("right", 3.18)))


def _configure_document_defaults(document: DocxDocument, config: dict[str, Any]):
    styles = document.styles
    normal = cast(Any, styles["Normal"])
    normal_font = config["document"]["default_font"]
    normal.font.name = str(normal_font.get("latin", "Times New Roman"))
    normal.font.size = Pt(float(normal_font.get("size", 12)))

    run_properties = normal._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:eastAsia"), str(normal_font.get("east_asia", "宋体")))

    _set_page_layout(document, config["document"]["margins_cm"])


def _extract_json_payload(raw_text: str) -> str:
    text = raw_text.strip()

    for prefix in _CODE_FENCE_PREFIXES:
        if text.lower().startswith(prefix):
            lines = text.splitlines()
            if len(lines) >= 2 and lines[-1].strip() == "```":
                text = "\n".join(lines[1:-1]).strip()
            break

    try:
        json.loads(text)
        return text
    except json.JSONDecodeError:
        pass

    start = text.find("[")
    end = text.rfind("]")
    if start != -1 and end != -1 and end > start:
        candidate = text[start:end + 1].strip()
        json.loads(candidate)
        return candidate

    raise json.JSONDecodeError("No JSON array found", text, 0)


def _parse_json_items(raw_text: str, required_keys: tuple[str, ...], name: str) -> list[dict[str, str]]:
    try:
        payload = _extract_json_payload(raw_text)
        data = json.loads(payload)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"{name} 不是合法 JSON 数组") from exc

    if not isinstance(data, list):
        raise RuntimeError(f"{name} 必须是 JSON 数组")

    parsed: list[dict[str, str]] = []
    for index, item in enumerate(data, start=1):
        if not isinstance(item, dict):
            raise RuntimeError(f"{name} 第 {index} 项不是 JSON 对象")

        normalized: dict[str, str] = {}
        for key in required_keys:
            normalized[key] = str(item.get(key, "")).strip()
        parsed.append(normalized)

    return parsed


def _add_metadata(document: DocxDocument, paper_path: str, discipline: str, generated_at: datetime, config: dict[str, Any]):
    metadata_lines = (
        ("论文文件", paper_path),
        ("学科领域", discipline),
        ("生成时间", generated_at.strftime("%Y-%m-%d %H:%M:%S")),
    )
    for label, value in metadata_lines:
        paragraph = document.add_paragraph()
        _set_paragraph_format(paragraph, config["paragraph"], first_line_chars=None)
        run = paragraph.add_run(f"{label}: {value}")
        _set_run_font(run, config["metadata"]["font"])


def _add_title(document: DocxDocument, title: str, config: dict[str, Any]):
    paragraph = document.add_paragraph()
    align = str(config["title"].get("align", "center")).lower()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(float(config["title"].get("space_after_pt", 12)))
    run = paragraph.add_run(title)
    _set_run_font(run, config["title"]["font"])


def _add_section_heading(document: DocxDocument, heading: str, config: dict[str, Any]):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(float(config["section_heading"].get("space_before_pt", 12)))
    paragraph.paragraph_format.space_after = Pt(float(config["section_heading"].get("space_after_pt", 6)))
    run = paragraph.add_run(heading)
    _set_run_font(run, config["section_heading"]["font"])


def _add_body_runs(paragraph, text: str, font_config: dict[str, Any]):
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(line)
        _set_run_font(run, font_config)


def _write_label_value_paragraph(
    document: DocxDocument,
    label: str,
    value: str,
    config: dict[str, Any],
    left_indent_cm: float,
):
    paragraph = document.add_paragraph()
    _set_paragraph_format(paragraph, config["paragraph"])
    paragraph.paragraph_format.left_indent = Cm(left_indent_cm)

    label_run = paragraph.add_run(label)
    _set_run_font(label_run, config["label_font"])

    if value:
        _add_body_runs(paragraph, value, config["body_font"])


def _write_review_items(document: DocxDocument, items: list[dict[str, str]], config: dict[str, Any]):
    review_config = config["review"]
    labels = review_config["labels"]
    left_indent = float(review_config.get("left_indent_cm", 0.74))

    for index, item in enumerate(items, start=1):
        # Keep number and issue label on the same line: "1. 问题：..."
        issue_paragraph = document.add_paragraph()
        _set_paragraph_format(issue_paragraph, config["paragraph"], first_line_chars=None)
        issue_paragraph.paragraph_format.left_indent = Cm(0)

        number_run = issue_paragraph.add_run(f"{index}. ")
        _set_run_font(number_run, review_config["number_font"])

        issue_label_run = issue_paragraph.add_run(str(labels.get("issue", "问题：")))
        _set_run_font(issue_label_run, review_config["label_font"])

        if item["issue"]:
            _add_body_runs(issue_paragraph, item["issue"], review_config["body_font"])

        location_paragraph = document.add_paragraph()
        _set_paragraph_format(location_paragraph, config["paragraph"], first_line_chars=None)
        location_paragraph.paragraph_format.left_indent = Cm(left_indent)
        marker = location_paragraph.add_run(str(labels.get("reference_location", "引用位置：")))
        _set_run_font(marker, review_config["label_font"])
        _set_rgb_color(marker, str(review_config.get("reference_color_rgb", "8B0000")))
        _add_body_runs(location_paragraph, item["reference location"], review_config["body_font"])

        _write_label_value_paragraph(
            document,
            str(labels.get("reference_text", "原文：")),
            item["reference text"],
            {
                "paragraph": config["paragraph"],
                "label_font": review_config["label_font"],
                "body_font": review_config["body_font"],
            },
            left_indent,
        )
        _write_label_value_paragraph(
            document,
            str(labels.get("detail", "详细描述：")),
            item["detail"],
            {
                "paragraph": config["paragraph"],
                "label_font": review_config["label_font"],
                "body_font": review_config["body_font"],
            },
            left_indent,
        )


def _write_response_items(document: DocxDocument, items: list[dict[str, str]], config: dict[str, Any]):
    response_config = config["response"]
    labels = response_config["labels"]
    left_indent = float(response_config.get("left_indent_cm", 0.74))

    for index, item in enumerate(items, start=1):
        title_paragraph = document.add_paragraph()
        _set_paragraph_format(title_paragraph, config["paragraph"], first_line_chars=None)
        title_run = title_paragraph.add_run(f"问题 {index}")
        _set_run_font(title_run, response_config["title_font"])

        _write_label_value_paragraph(
            document,
            str(labels.get("problem", "审稿问题总结：")),
            item["problem"],
            {
                "paragraph": config["paragraph"],
                "label_font": response_config["label_font"],
                "body_font": response_config["body_font"],
            },
            left_indent,
        )
        _write_label_value_paragraph(
            document,
            str(labels.get("responde", "回应与解决方法：")),
            item["responde"],
            {
                "paragraph": config["paragraph"],
                "label_font": response_config["label_font"],
                "body_font": response_config["body_font"],
            },
            left_indent,
        )


def _add_footer_note(document: DocxDocument, note: str, config: dict[str, Any]):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(float(config["footer"].get("space_before_pt", 18)))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(note)
    _set_run_font(run, config["footer"]["font"])
    _set_rgb_color(run, str(config["footer"].get("color_rgb", "666666")))


def export_review_documents(
    output_dir: str,
    paper_path: str,
    discipline: str,
    review_questions: str,
    review_responses: str,
    generated_at: datetime | None = None,
    format_path: str | None = None,
) -> tuple[str, str]:
    """Export review questions and responses to two docx files."""
    try:
        os.makedirs(output_dir, exist_ok=True)
    except OSError as exc:
        raise RuntimeError(f"无法创建输出目录: {output_dir}") from exc

    config = _load_format_config(format_path)
    review_items = _parse_json_items(review_questions, _REVIEW_REQUIRED_KEYS, "reviewer 输出")
    response_items = _parse_json_items(review_responses, _RESPONSE_REQUIRED_KEYS, "responder 输出")

    now = generated_at or datetime.now()
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    review_path = os.path.join(output_dir, f"评审意见+{timestamp}.docx")
    response_path = os.path.join(output_dir, f"回复+{timestamp}.docx")

    review_doc = Document()
    _configure_document_defaults(review_doc, config)
    _add_title(review_doc, "学术论文评审意见", config)
    _add_metadata(review_doc, paper_path, discipline, now, config)
    _add_section_heading(review_doc, "评审意见正文", config)
    _write_review_items(review_doc, review_items, config)
    _add_footer_note(review_doc, "注：引用定位内容直接保留原文标注，便于与论文正文逐项核对。", config)
    review_doc.save(review_path)

    response_doc = Document()
    _configure_document_defaults(response_doc, config)
    _add_title(response_doc, "学术论文评审回复", config)
    _add_metadata(response_doc, paper_path, discipline, now, config)
    _add_section_heading(response_doc, "作者回复正文", config)
    _write_response_items(response_doc, response_items, config)
    _add_footer_note(response_doc, "注：回复按问题编号组织，建议结合修订稿逐条对应落实。", config)
    response_doc.save(response_path)

    return review_path, response_path
